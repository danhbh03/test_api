import os
import gradio as gr
import asyncio
from docx import Document
from lightrag import LightRAG, QueryParam
from lightrag.llm.ollama import ollama_embed
from lightrag.kg.shared_storage import initialize_pipeline_status
from lightrag.utils import EmbeddingFunc, setup_logger
import subprocess
from transformers import AutoTokenizer
import re
import unicodedata
from groq import Groq
from dotenv import load_dotenv

# ==========================
# Setup & Globals
# ==========================
setup_logger("lightrag", level="INFO")

WORKING_DIR = os.path.abspath("./test_chat_rag_storage")
os.makedirs(WORKING_DIR, exist_ok=True)
os.environ["LIGHTRAG_WORKING_DIR"] = WORKING_DIR
print("📂 Danh sách file sau khi xóa:", os.listdir(WORKING_DIR))

TOKENIZER_MODEL = "gpt2"
tokenizer = AutoTokenizer.from_pretrained(TOKENIZER_MODEL)
load_dotenv()

# Khởi tạo client với API key từ .env
client = Groq(api_key=os.getenv("GROQ_API_KEY"))

# ==========================
# Hàm LLM cho LightRAG (đã sửa)
# ==========================
async def groq_model_complete(
    prompt: str,
    system_prompt: str | None = None,
    history_messages: list[dict] | None = None,
    **kwargs
) -> str:
    try:
        messages = []
        if system_prompt:
            messages.append({"role": "system", "content": system_prompt})
        if history_messages:
            messages.extend(history_messages)
        messages.append({"role": "user", "content": prompt})

        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=messages,
            temperature=kwargs.get("temperature", 0.7),
            top_p=kwargs.get("top_p", 1),
            max_tokens=kwargs.get("max_tokens", 1024),
            stream=False,
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        raise RuntimeError(f"[ERROR from Groq LLM] {e}") from e

# ==========================
# Helpers
# ==========================
def fetch_uploaded_gradio_file(file) -> str:
    try:
        if file.name.endswith(".docx"):
            doc = Document(file.name)
            full_text = [para.text for para in doc.paragraphs if para.text.strip()]
            return '\n'.join(full_text)
        with open(file.name, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        raise Exception(f"Lỗi đọc file: {e}")

def num_tokens(text: str) -> int:
    return len(tokenizer.encode(text, add_special_tokens=False))

def sentence_split(text: str):
    sentences = re.split(r'(?<=[.!?。！？])\s+', text)
    return [s.strip() for s in sentences if s.strip()]

def pre_split(text: str, max_block_tokens: int = 2000):
    sentences = sentence_split(text)
    blocks, current_block, current_tokens = [], [], 0
    for sent in sentences:
        sent_tokens = tokenizer.encode(sent, add_special_tokens=False)
        if current_tokens + len(sent_tokens) > max_block_tokens:
            blocks.append(" ".join(current_block))
            current_block, current_tokens = [], 0
        current_block.append(sent)
        current_tokens += len(sent_tokens)
    if current_block:
        blocks.append(" ".join(current_block))
    return blocks

def enforce_max_len(chunk: str, max_tokens: int = 400, overlap: int = 50):
    tokens = tokenizer.encode(chunk, add_special_tokens=False)
    if len(tokens) <= max_tokens:
        return [chunk]
    final_chunks, step = [], max_tokens - max(0, overlap)
    for i in range(0, len(tokens), step):
        sub = tokens[i:i + max_tokens]
        final_chunks.append(tokenizer.decode(sub))
    return final_chunks

def split_paragraph_semantic(paragraph: str, max_tokens: int = 500, overlap: int = 100):
    # fallback splitter, tránh ollama subprocess nếu không có model
    blocks = pre_split(paragraph, max_block_tokens=1800)
    all_chunks = []
    for block in blocks:
        all_chunks.extend(enforce_max_len(block, max_tokens=max_tokens, overlap=overlap))
    return all_chunks

def prepare_document(file):
    """Đọc file txt/docx + chuẩn hóa + tách đoạn"""
    try:
        if file is None:
            return []
        if file.name.endswith(".docx"):
            doc = Document(file.name)
            raw_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        else:
            with open(file.name, "r", encoding="utf-8") as f:
                raw_text = f.read()
        text = unicodedata.normalize("NFC", raw_text)
        text = text.replace("\t", " ")
        text = re.sub(r"[ ]{2,}", " ", text)
        text = re.sub(r"\r\n", "\n", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = text.strip()
        paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
        seen, out = set(), []
        for p in paragraphs:
            if p not in seen:
                seen.add(p)
                out.append(p)
        return out
    except Exception as e:
        raise Exception(f"Lỗi đọc file: {e}")

# ==========================
# LightRAG init
# ==========================
async def initialize_chatbot_rag():
    rag = LightRAG(
        working_dir=WORKING_DIR,
        llm_model_func=groq_model_complete,
        llm_model_name="llama-3.3-70b-versatile",
        embedding_func=EmbeddingFunc(
            embedding_dim=1024,
            max_token_size=8192,
            func=lambda texts: ollama_embed(texts, embed_model="bge-m3"),
        )
    )
    await rag.initialize_storages()
    await initialize_pipeline_status()
    return rag

# ==========================
# Ingestion
# ==========================
async def add_document_to_rag(file, rag: LightRAG):
    try:
        paragraphs = prepare_document(file)
        if not paragraphs:
            return "⚠️ File rỗng hoặc không đọc được nội dung.", False
        success_count, failed_items = 0, []
        for i, paragraph in enumerate(paragraphs, 1):
            try:
                sub_chunks = split_paragraph_semantic(paragraph, max_tokens=400, overlap=50)
                for sub in sub_chunks:
                    sub = sub.strip()
                    if not sub:
                        continue
                    await rag.ainsert(sub)
                    success_count += 1
                    print(f"✅ Đã thêm chunk {success_count}")
            except Exception as line_error:
                failed_items.append(f"Đoạn {i}: {line_error}")
                print(f"❌ Lỗi đoạn {i}: {line_error}")
        if failed_items:
            msg = f"⚠️ Thêm {success_count} chunk thành công. {len(failed_items)} mục lỗi:\n" + "\n".join(failed_items[:5])
            return msg, False
        else:
            return f"✅ Đã thêm {success_count} chunk tài liệu thành công!", True
    except Exception as e:
        return f"❌ Lỗi khi xử lý file: {e}", False

# ==========================
# Query / Chat
# ==========================
RESPONSE_RULES = (
    "You must strictly extract information from the inserted documents."
    "Do not add personal opinions or fabricate any facts."
    "You may paraphrase and summarize only if the meaning is preserved."
    "Do not include metadata such as References, Created date, Entity name, or 'from the Knowledge Graph' in the response."
    "If multiple retrieved passages have similar meaning, keep only the earliest one."
)

def to_light_history(hist):
    out = []
    for m in hist:
        role = m.get("role")
        content = m.get("content", "")
        if not content:
            continue
        content = content.split("\n\n📊")[0]
        out.append({"role": role, "content": content})
    return out

async def RAG_chatbot(message, history, rag, file, query_mode):
    history = history or []
    prompt = f"{RESPONSE_RULES}\n\nTopic: {message}"

    qp = QueryParam(
        mode=query_mode,   # <-- local hoặc hybrid
        top_k=10,
        user_prompt=prompt,
        conversation_history=to_light_history(history)
    )

    response = await rag.aquery(message, qp)

    history.append({"role": "user", "content": message})
    history.append({"role": "assistant", "content": response})
    return history, history

def create_rag_wrapper(rag):
    async def rag_wrapper(message, history, file, query_mode):
        return await RAG_chatbot(message, history, rag, file, query_mode)
    return rag_wrapper

# ==========================
# Gradio UI
# ==========================
async def main():
    rag = await initialize_chatbot_rag()
    with gr.Blocks() as demo:
        with gr.Row():
            with gr.Column(scale=4):
                gr.Markdown("## 🧠 Chatbot tài liệu dùng LightRAG (Chuẩn hóa + Semantic Chunking)")
            with gr.Column(scale=1):
                query_mode = gr.Dropdown(
                    choices=["local", "hybrid"],
                    value="hybrid",
                    label="⚙️ Chế độ",
                )

        rag_state = gr.State(rag)
        is_rag_ready = gr.State(False)
        chatbot_ui = gr.Chatbot(type="messages")
        msg = gr.Textbox(label="💬 Nhập câu hỏi")
        history = gr.State([])

        with gr.Row():
            send_btn = gr.Button("📩 Gửi")
            upload_btn = gr.Button("📥 Thêm tài liệu")

        file_input = gr.File(label="📄 Upload file .txt/.docx", file_types=[".txt", ".docx"])
        upload_result = gr.Textbox(label="Kết quả thêm tài liệu", interactive=False)

        upload_btn.click(
            fn=add_document_to_rag,
            inputs=[file_input, rag_state],
            outputs=[upload_result, is_rag_ready]
        )

        rag_wrapper_func = create_rag_wrapper(rag)

        send_btn.click(
            fn=rag_wrapper_func,
            inputs=[msg, history, file_input, query_mode],
            outputs=[chatbot_ui, history]
        )

        msg.submit(
            fn=rag_wrapper_func,
            inputs=[msg, history, file_input, query_mode],
            outputs=[chatbot_ui, history]
        )

        demo.launch(server_name="127.0.0.1", server_port=9621, show_api=False, share=True)

if __name__ == "__main__":
    import nest_asyncio
    nest_asyncio.apply()
    asyncio.run(main())
